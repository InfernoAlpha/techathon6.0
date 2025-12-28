import os
from typing import List, Union
from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from langchain.tools import tool
from schema import doc_gen_format

@tool
def doc_specific_gen(data:doc_gen_format) -> dict:
    """
    this is a tool used to convert text into a word docx file given the contents for the file
    """

    doc = Document()

    styles = doc.styles
    if 'MyCalibriStyle' not in styles:
        new_style = styles.add_style('MyCalibriStyle', WD_STYLE_TYPE.PARAGRAPH)
        new_style.font.name = 'Calibri'
        new_style.font.size = Pt(12)
        new_style.font._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')

    if 'BoldCalibriStyle' not in styles:
        bold_style = styles.add_style('BoldCalibriStyle', WD_STYLE_TYPE.PARAGRAPH)
        bold_style.font.name = 'Calibri'
        bold_style.font.size = Pt(12)
        bold_style.font.bold = True
        bold_style.font._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')

    para = doc.add_paragraph(f"                                     RCA/CAPA REPORT                                          ")
    para.style = 'BoldCalibriStyle'

    para = doc.add_paragraph(data.model_dump(by_alias=True)['text'])
    para.style = 'BoldCalibriStyle'

    doc.save("RCA_CAPA_report.docx")
    print("Document generation complete.")
    return "docx file has been created"

if(__name__ == "__main__"):
    doc_specific_gen("sample testig","sample.docx")